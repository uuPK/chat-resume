import os
import uuid
from typing import Dict, Any
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import settings


class ExportService:
    def __init__(self):
        self.export_dir = os.path.join(settings.UPLOAD_DIR, "exports")
        os.makedirs(self.export_dir, exist_ok=True)

        # 注册中文字体（如果有的话）
        try:
            # 这里可以添加中文字体支持
            # pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
            pass
        except Exception:
            pass

    def export_to_pdf(
        self, resume_content: Dict[str, Any], template: str = "default"
    ) -> str:
        """导出简历为PDF"""

        # 生成唯一文件名
        filename = f"resume_{uuid.uuid4().hex}.pdf"
        filepath = os.path.join(self.export_dir, filename)

        # 创建PDF文档
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []

        # 获取样式
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=colors.black,
            spaceAfter=12,
            alignment=1,  # 居中
        )

        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=6,
            spaceBefore=12,
        )

        normal_style = styles["Normal"]

        # 添加个人信息
        personal_info = resume_content.get("personal_info", {})
        if personal_info.get("name"):
            story.append(Paragraph(personal_info["name"], title_style))
            story.append(Spacer(1, 12))

        # 添加联系信息
        contact_info = []
        if personal_info.get("email"):
            contact_info.append(f"邮箱: {personal_info['email']}")
        if personal_info.get("phone"):
            contact_info.append(f"电话: {personal_info['phone']}")
        if personal_info.get("address"):
            contact_info.append(f"地址: {personal_info['address']}")

        if contact_info:
            story.append(Paragraph(" | ".join(contact_info), normal_style))
            story.append(Spacer(1, 12))

        # 添加教育背景
        education = resume_content.get("education", [])
        if education:
            story.append(Paragraph("教育背景", heading_style))
            for edu in education:
                edu_text = f"<b>{edu.get('school', '')}</b>"
                if edu.get("degree"):
                    edu_text += f" - {edu.get('degree')}"
                if edu.get("major"):
                    edu_text += f" - {edu.get('major')}"
                if edu.get("duration"):
                    edu_text += f" ({edu.get('duration')})"
                story.append(Paragraph(edu_text, normal_style))
            story.append(Spacer(1, 12))

        # 添加工作经验
        work_experience = resume_content.get("work_experience", [])
        if work_experience:
            story.append(Paragraph("工作经验", heading_style))
            for work in work_experience:
                company_text = f"<b>{work.get('company', '')}</b>"
                if work.get("position"):
                    company_text += f" - {work.get('position')}"
                if work.get("duration"):
                    company_text += f" ({work.get('duration')})"
                story.append(Paragraph(company_text, normal_style))

                for resp in work.get("responsibilities", []):
                    story.append(Paragraph(f"• {resp}", normal_style))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))

        # 添加技能
        skills = resume_content.get("skills", [])
        if skills:
            story.append(Paragraph("技能", heading_style))
            skills_text = " | ".join(skills)
            story.append(Paragraph(skills_text, normal_style))
            story.append(Spacer(1, 12))

        # 添加项目经验
        projects = resume_content.get("projects", [])
        if projects:
            story.append(Paragraph("项目经验", heading_style))
            for project in projects:
                project_text = f"<b>{project.get('name', '')}</b>"
                story.append(Paragraph(project_text, normal_style))
                story.append(Paragraph(project.get("description", ""), normal_style))
                story.append(Spacer(1, 6))

        # 构建PDF
        doc.build(story)

        return filepath

    def export_to_docx(
        self, resume_content: Dict[str, Any], template: str = "default"
    ) -> str:
        """导出简历为Word文档"""

        # 生成唯一文件名
        filename = f"resume_{uuid.uuid4().hex}.docx"
        filepath = os.path.join(self.export_dir, filename)

        # 创建Word文档
        doc = Document()

        # 添加个人信息
        personal_info = resume_content.get("personal_info", {})
        if personal_info.get("name"):
            title = doc.add_heading(personal_info["name"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加联系信息
        contact_info = []
        if personal_info.get("email"):
            contact_info.append(f"邮箱: {personal_info['email']}")
        if personal_info.get("phone"):
            contact_info.append(f"电话: {personal_info['phone']}")
        if personal_info.get("address"):
            contact_info.append(f"地址: {personal_info['address']}")

        if contact_info:
            contact_para = doc.add_paragraph(" | ".join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加教育背景
        education = resume_content.get("education", [])
        if education:
            doc.add_heading("教育背景", level=1)
            for edu in education:
                edu_text = edu.get("school", "")
                if edu.get("degree"):
                    edu_text += f" - {edu.get('degree')}"
                if edu.get("major"):
                    edu_text += f" - {edu.get('major')}"
                if edu.get("duration"):
                    edu_text += f" ({edu.get('duration')})"
                doc.add_paragraph(edu_text, style="List Bullet")

        # 添加工作经验
        work_experience = resume_content.get("work_experience", [])
        if work_experience:
            doc.add_heading("工作经验", level=1)
            for work in work_experience:
                company_text = work.get("company", "")
                if work.get("position"):
                    company_text += f" - {work.get('position')}"
                if work.get("duration"):
                    company_text += f" ({work.get('duration')})"
                doc.add_paragraph(company_text, style="List Bullet")

                for resp in work.get("responsibilities", []):
                    doc.add_paragraph(f"• {resp}", style="List Bullet 2")

        # 添加技能
        skills = resume_content.get("skills", [])
        if skills:
            doc.add_heading("技能", level=1)
            skills_text = " | ".join(skills)
            doc.add_paragraph(skills_text)

        # 添加项目经验
        projects = resume_content.get("projects", [])
        if projects:
            doc.add_heading("项目经验", level=1)
            for project in projects:
                doc.add_paragraph(project.get("name", ""), style="List Bullet")
                doc.add_paragraph(project.get("description", ""), style="List Bullet 2")

        # 保存文档
        doc.save(filepath)

        return filepath

    def export_to_html(
        self, resume_content: Dict[str, Any], template: str = "default"
    ) -> str:
        """导出简历为HTML"""

        # 生成唯一文件名
        filename = f"resume_{uuid.uuid4().hex}.html"
        filepath = os.path.join(self.export_dir, filename)

        # 构建HTML内容
        html_content = self._build_html_content(resume_content, template)

        # 保存HTML文件
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html_content)

        return filepath

    def _build_html_content(self, resume_content: Dict[str, Any], template: str) -> str:
        """构建HTML内容"""

        personal_info = resume_content.get("personal_info", {})
        education = resume_content.get("education", [])
        work_experience = resume_content.get("work_experience", [])
        skills = resume_content.get("skills", [])
        projects = resume_content.get("projects", [])

        html = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>简历 - {personal_info.get("name", "")}</title>
            <style>
                body {{
                    font-family: 'Microsoft YaHei', Arial, sans-serif;
                    line-height: 1.6;
                    margin: 0;
                    padding: 20px;
                    background-color: #f5f5f5;
                }}
                .container {{
                    max-width: 800px;
                    margin: 0 auto;
                    background-color: white;
                    padding: 40px;
                    border-radius: 10px;
                    box-shadow: 0 0 10px rgba(0,0,0,0.1);
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                    border-bottom: 2px solid #007bff;
                    padding-bottom: 20px;
                }}
                .name {{
                    font-size: 2.5em;
                    font-weight: bold;
                    color: #333;
                    margin-bottom: 10px;
                }}
                .contact {{
                    font-size: 1.1em;
                    color: #666;
                }}
                .section {{
                    margin-bottom: 30px;
                }}
                .section-title {{
                    font-size: 1.5em;
                    font-weight: bold;
                    color: #007bff;
                    border-bottom: 1px solid #eee;
                    padding-bottom: 5px;
                    margin-bottom: 15px;
                }}
                .item {{
                    margin-bottom: 15px;
                }}
                .item-title {{
                    font-weight: bold;
                    color: #333;
                }}
                .item-subtitle {{
                    color: #666;
                    font-style: italic;
                }}
                .item-content {{
                    margin-top: 5px;
                    color: #555;
                }}
                .skills {{
                    display: flex;
                    flex-wrap: wrap;
                    gap: 10px;
                }}
                .skill {{
                    background-color: #e9ecef;
                    padding: 5px 10px;
                    border-radius: 15px;
                    font-size: 0.9em;
                }}
                ul {{
                    padding-left: 20px;
                }}
                li {{
                    margin-bottom: 5px;
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <div class="name">{personal_info.get("name", "")}</div>
                    <div class="contact">
        """

        # 添加联系信息
        contact_info = []
        if personal_info.get("email"):
            contact_info.append(f"邮箱: {personal_info['email']}")
        if personal_info.get("phone"):
            contact_info.append(f"电话: {personal_info['phone']}")
        if personal_info.get("address"):
            contact_info.append(f"地址: {personal_info['address']}")

        html += " | ".join(contact_info)
        html += """
                    </div>
                </div>
        """

        # 添加教育背景
        if education:
            html += """
                <div class="section">
                    <div class="section-title">教育背景</div>
            """
            for edu in education:
                html += f"""
                    <div class="item">
                        <div class="item-title">{edu.get("school", "")}</div>
                        <div class="item-subtitle">{edu.get("degree", "")} - {edu.get("major", "")} ({edu.get("duration", "")})</div>
                    </div>
                """
            html += "</div>"

        # 添加工作经验
        if work_experience:
            html += """
                <div class="section">
                    <div class="section-title">工作经验</div>
            """
            for work in work_experience:
                html += f"""
                    <div class="item">
                        <div class="item-title">{work.get("company", "")}</div>
                        <div class="item-subtitle">{work.get("position", "")} ({work.get("duration", "")})</div>
                        <div class="item-content">
                            <ul>
                """
                for resp in work.get("responsibilities", []):
                    html += f"<li>{resp}</li>"
                html += """
                            </ul>
                        </div>
                    </div>
                """
            html += "</div>"

        # 添加技能
        if skills:
            html += """
                <div class="section">
                    <div class="section-title">技能</div>
                    <div class="skills">
            """
            for skill in skills:
                html += f'<span class="skill">{skill}</span>'
            html += """
                    </div>
                </div>
            """

        # 添加项目经验
        if projects:
            html += """
                <div class="section">
                    <div class="section-title">项目经验</div>
            """
            for project in projects:
                html += f"""
                    <div class="item">
                        <div class="item-title">{project.get("name", "")}</div>
                        <div class="item-content">{project.get("description", "")}</div>
                    </div>
                """
            html += "</div>"

        html += """
            </div>
        </body>
        </html>
        """

        return html

    def get_file_url(self, filepath: str) -> str:
        """获取文件的访问URL"""
        filename = os.path.basename(filepath)
        return f"/api/v1/export/download/{filename}"

    def delete_file(self, filepath: str) -> bool:
        """删除导出的文件"""
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                return True
            return False
        except Exception:
            return False
